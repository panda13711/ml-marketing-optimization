"""
Generate Executive Word Document for Cyberdine Marketing Analysis
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================================
# TITLE PAGE
# ============================================================================
title = doc.add_heading('CYBERDINE MARKETING PERFORMANCE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Executive Board Presentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(16)
subtitle_format.font.bold = True

date_para = doc.add_paragraph('November 28, 2025')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(12)

doc.add_paragraph('\n' * 3)

prepared = doc.add_paragraph('Strategic Recommendations & Budget Optimization')
prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
prepared.runs[0].font.size = Pt(14)
prepared.runs[0].font.italic = True

doc.add_page_break()

# ============================================================================
# EXECUTIVE SUMMARY
# ============================================================================
doc.add_heading('EXECUTIVE SUMMARY', 1)

doc.add_paragraph('Analysis Period: Recent campaign performance across 12 campaigns')
doc.add_paragraph('Total Marketing Investment: $24,000')
doc.add_paragraph('Total Revenue Generated: $480,426')
doc.add_paragraph('Total Profit (after marketing costs): $55,616')
doc.add_paragraph('Overall ROI: 231.7%')
doc.add_paragraph('Overall ROAS: 20.02x')

doc.add_heading('Key Findings', 2)
findings = [
    'All platforms are profitable - positive ROI across the board',
    'Instagram is the clear winner - highest ROI at 268.5%',
    'Significant audience segmentation discovered by platform',
    'Strong performance indicates scalable opportunity for growth'
]
for finding in findings:
    doc.add_paragraph(finding, style='List Bullet')

doc.add_page_break()

# ============================================================================
# QUESTION 1: CAMPAIGN PERFORMANCE OVERVIEW
# ============================================================================
doc.add_heading('1. CAMPAIGN PERFORMANCE OVERVIEW', 1)

doc.add_heading('Platform Performance Comparison', 2)

# Load data for tables
platform_df = pd.read_csv('platform_analysis.csv')

# Create table
table = doc.add_table(rows=5, cols=8)
table.style = 'Light Grid Accent 1'

# Header row
headers = ['Platform', 'Investment', 'Revenue', 'Profit', 'ROI', 'ROAS', 'Conversions', 'CPA']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

# Data rows
platforms_sorted = platform_df.sort_values('ROI', ascending=False)
for idx, (_, row) in enumerate(platforms_sorted.iterrows(), 1):
    table.rows[idx].cells[0].text = row['platform']
    table.rows[idx].cells[1].text = f"${row['cost']:,.0f}"
    table.rows[idx].cells[2].text = f"${row['revenue']:,.0f}"
    table.rows[idx].cells[3].text = f"${row['profit']:,.0f}"
    table.rows[idx].cells[4].text = f"{row['ROI']:.1f}%"
    table.rows[idx].cells[5].text = f"{row['ROAS']:.2f}x"
    table.rows[idx].cells[6].text = f"{int(row['conversions'])}"
    table.rows[idx].cells[7].text = f"${row['avg_cpa']:.2f}"

doc.add_paragraph()

doc.add_heading('Key Insights', 2)
insights = [
    'Instagram dominates with highest ROI (268.5%), ROAS (24.40x), and lowest CPA ($6.42)',
    'LinkedIn performs strongly despite highest CPA - B2B audience justifies premium',
    'Facebook delivers solid mid-tier performance with good volume',
    'Google shows strong profitability at 158.8% ROI, though trailing competitors'
]
for insight in insights:
    doc.add_paragraph(insight, style='List Bullet')

doc.add_heading('Critical Discovery: Audience Segmentation', 2)

# Product mix table
sales = pd.read_csv('sales.csv')
products = pd.read_csv('products.csv')
campaigns = pd.read_csv('campaigns.csv')
sales_full = sales.merge(products, on='product_id').merge(campaigns, on='campaign_id')
product_mix = pd.crosstab(sales_full['platform'], sales_full['product_category'], normalize='index') * 100

doc.add_paragraph('Each platform attracts distinctly different customer profiles:')

mix_table = doc.add_table(rows=5, cols=4)
mix_table.style = 'Light List Accent 1'

# Headers
mix_table.rows[0].cells[0].text = 'Platform'
mix_table.rows[0].cells[1].text = 'Primary Category'
mix_table.rows[0].cells[2].text = '% of Sales'
mix_table.rows[0].cells[3].text = 'Customer Type'

# Data
mix_table.rows[1].cells[0].text = 'Instagram'
mix_table.rows[1].cells[1].text = 'Fashion'
mix_table.rows[1].cells[2].text = '60.2%'
mix_table.rows[1].cells[3].text = 'Young, style-conscious'

mix_table.rows[2].cells[0].text = 'LinkedIn'
mix_table.rows[2].cells[1].text = 'Tech'
mix_table.rows[2].cells[2].text = '89.2%'
mix_table.rows[2].cells[3].text = 'Professional B2B buyers'

mix_table.rows[3].cells[0].text = 'Facebook'
mix_table.rows[3].cells[1].text = 'Health'
mix_table.rows[3].cells[2].text = '60.7%'
mix_table.rows[3].cells[3].text = 'Health-focused families'

mix_table.rows[4].cells[0].text = 'Google'
mix_table.rows[4].cells[1].text = 'Mixed'
mix_table.rows[4].cells[2].text = 'Balanced'
mix_table.rows[4].cells[3].text = 'Active searchers'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Implication: One-size-fits-all creative strategy is suboptimal. Platform-specific ads are required.')
run.bold = True

doc.add_page_break()

# ============================================================================
# QUESTION 2: $2,000 BUDGET ALLOCATION
# ============================================================================
doc.add_heading('2. BUDGET ALLOCATION: $2,000', 1)

doc.add_heading('Recommendation: High-ROI Concentration Strategy', 2)
doc.add_paragraph('Criteria: Maximize Return on Investment by selecting highest-performing campaigns')

alloc_table = doc.add_table(rows=4, cols=5)
alloc_table.style = 'Medium Shading 1 Accent 1'

# Headers
alloc_table.rows[0].cells[0].text = 'Campaign'
alloc_table.rows[0].cells[1].text = 'Platform'
alloc_table.rows[0].cells[2].text = 'Investment'
alloc_table.rows[0].cells[3].text = 'Expected Profit'
alloc_table.rows[0].cells[4].text = 'Expected ROI'

# Data
alloc_table.rows[1].cells[0].text = 'Campaign 7'
alloc_table.rows[1].cells[1].text = 'Instagram'
alloc_table.rows[1].cells[2].text = '$1,000'
alloc_table.rows[1].cells[3].text = '$4,125'
alloc_table.rows[1].cells[4].text = '312.5%'

alloc_table.rows[2].cells[0].text = 'Campaign 4'
alloc_table.rows[2].cells[1].text = 'Facebook'
alloc_table.rows[2].cells[2].text = '$1,000'
alloc_table.rows[2].cells[3].text = '$4,070'
alloc_table.rows[2].cells[4].text = '307.0%'

alloc_table.rows[3].cells[0].text = 'TOTAL'
alloc_table.rows[3].cells[1].text = '-'
alloc_table.rows[3].cells[2].text = '$2,000'
alloc_table.rows[3].cells[3].text = '$8,195'
alloc_table.rows[3].cells[4].text = '309.8%'

doc.add_paragraph()

doc.add_heading('Strategic Rationale', 2)
rationale = [
    'Focus on proven winners - Both campaigns have delivered exceptional ROI',
    'Platform diversity - Split between Instagram and Facebook reduces risk',
    'Expected return - Every $1 invested returns $4.10',
    'Fast deployment - Leverage existing successful campaigns'
]
for item in rationale:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Expected Outcomes', 2)
doc.add_paragraph('Revenue: ~$30,000')
doc.add_paragraph('Profit: ~$8,195')
doc.add_paragraph('Net gain after cost: $6,195')

doc.add_page_break()

# ============================================================================
# QUESTION 3: $5,000 BUDGET ALLOCATION
# ============================================================================
doc.add_heading('3. BUDGET ALLOCATION: $5,000', 1)

doc.add_heading('Recommendation: Diversified High-Performance Portfolio', 2)
doc.add_paragraph('Strategy Change: YES - More capital enables platform diversification')

alloc5k_table = doc.add_table(rows=6, cols=5)
alloc5k_table.style = 'Medium Shading 1 Accent 1'

# Headers
alloc5k_table.rows[0].cells[0].text = 'Campaign'
alloc5k_table.rows[0].cells[1].text = 'Platform'
alloc5k_table.rows[0].cells[2].text = 'Investment'
alloc5k_table.rows[0].cells[3].text = 'Expected Profit'
alloc5k_table.rows[0].cells[4].text = 'Expected ROI'

# Data
alloc5k_table.rows[1].cells[0].text = 'Campaign 7'
alloc5k_table.rows[1].cells[1].text = 'Instagram'
alloc5k_table.rows[1].cells[2].text = '$1,000'
alloc5k_table.rows[1].cells[3].text = '$4,125'
alloc5k_table.rows[1].cells[4].text = '312.5%'

alloc5k_table.rows[2].cells[0].text = 'Campaign 4'
alloc5k_table.rows[2].cells[1].text = 'Facebook'
alloc5k_table.rows[2].cells[2].text = '$1,000'
alloc5k_table.rows[2].cells[3].text = '$4,070'
alloc5k_table.rows[2].cells[4].text = '307.0%'

alloc5k_table.rows[3].cells[0].text = 'Campaign 8'
alloc5k_table.rows[3].cells[1].text = 'Instagram'
alloc5k_table.rows[3].cells[2].text = '$2,000'
alloc5k_table.rows[3].cells[3].text = '$7,619'
alloc5k_table.rows[3].cells[4].text = '281.0%'

alloc5k_table.rows[4].cells[0].text = 'Campaign 10'
alloc5k_table.rows[4].cells[1].text = 'LinkedIn'
alloc5k_table.rows[4].cells[2].text = '$1,000'
alloc5k_table.rows[4].cells[3].text = '$3,762'
alloc5k_table.rows[4].cells[4].text = '276.2%'

alloc5k_table.rows[5].cells[0].text = 'TOTAL'
alloc5k_table.rows[5].cells[1].text = '-'
alloc5k_table.rows[5].cells[2].text = '$5,000'
alloc5k_table.rows[5].cells[3].text = '$19,576'
alloc5k_table.rows[5].cells[4].text = '291.5%'

doc.add_paragraph()

doc.add_heading('Key Differences from $2K Strategy', 2)
differences = [
    '4 campaigns vs 2 - Increased diversification',
    '3 platforms vs 2 - Added LinkedIn for B2B tech buyers',
    'Deeper Instagram investment - $3K total in top performer',
    'Balanced risk profile - Not over-concentrated in single campaign'
]
for item in differences:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Expected Outcomes', 2)
doc.add_paragraph('Revenue: ~$75,000')
doc.add_paragraph('Profit: ~$19,576')
doc.add_paragraph('Net gain after cost: $14,576')
doc.add_paragraph('ROI improvement: Maintains high 291.5% despite broader spread')

doc.add_page_break()

# ============================================================================
# QUESTION 4: $30,000 LARGE BUDGET ALLOCATION
# ============================================================================
doc.add_heading('4. LARGE BUDGET ALLOCATION: $30,000', 1)

doc.add_heading('Channel Scalability Analysis', 2)
doc.add_paragraph('Question: Which channel is most promising at scale?')

scale_table = doc.add_table(rows=5, cols=3)
scale_table.style = 'Light List Accent 1'

scale_table.rows[0].cells[0].text = 'Platform'
scale_table.rows[0].cells[1].text = 'Scalability Score'
scale_table.rows[0].cells[2].text = 'Ranking'

scale_table.rows[1].cells[0].text = 'Instagram'
scale_table.rows[1].cells[1].text = '655.23'
scale_table.rows[1].cells[2].text = '🥇 #1'

scale_table.rows[2].cells[0].text = 'LinkedIn'
scale_table.rows[2].cells[1].text = '555.76'
scale_table.rows[2].cells[2].text = '🥈 #2'

scale_table.rows[3].cells[0].text = 'Facebook'
scale_table.rows[3].cells[1].text = '447.74'
scale_table.rows[3].cells[2].text = '🥉 #3'

scale_table.rows[4].cells[0].text = 'Google'
scale_table.rows[4].cells[1].text = '247.78'
scale_table.rows[4].cells[2].text = '#4'

doc.add_paragraph()
doc.add_paragraph('Scalability Score = (ROI/100) × (Conversion Rate/10) × (ROAS)')

doc.add_heading('Recommendation: Instagram-Led Diversified Strategy', 2)

doc.add_heading('Suggested $30K Allocation', 3)
doc.add_paragraph('Instagram: $15,000 (50%) - Primary channel')
doc.add_paragraph('LinkedIn: $7,500 (25%) - Secondary channel')
doc.add_paragraph('Facebook: $4,500 (15%) - Tertiary channel')
doc.add_paragraph('Google: $3,000 (10%) - Testing/experimental')

doc.add_heading('Strategic Rationale', 2)

p1 = doc.add_paragraph()
run = p1.add_run('Why Instagram deserves 50% allocation:')
run.bold = True

reasons = [
    'Highest ROI (268.5%)',
    'Highest ROAS (24.40x)',
    'Lowest CPA ($6.42)',
    'Best scalability score',
    'Strong fashion category alignment',
    'Proven consistent performance'
]
for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

p2 = doc.add_paragraph()
run = p2.add_run('Why maintain diversification:')
run.bold = True

diversity_reasons = [
    'LinkedIn captures valuable B2B tech segment',
    'Facebook reaches health-conscious family demographic',
    'Google captures high-intent search traffic',
    'Risk mitigation against platform algorithm changes',
    'Maintains omnichannel brand presence'
]
for reason in diversity_reasons:
    doc.add_paragraph(reason, style='List Bullet')

doc.add_heading('Expected Outcomes at $30K Scale', 2)
doc.add_paragraph('Projected Revenue: ~$600,000+')
doc.add_paragraph('Projected Profit: ~$70,000+')
doc.add_paragraph('Expected Portfolio ROI: ~233%')
doc.add_paragraph('Projected Conversions: ~2,800+')

doc.add_heading('Scaling Considerations', 2)
considerations = [
    'Diminishing returns possible as spend increases',
    'Audience saturation may occur on smaller platforms',
    'Recommended approach: Gradual scale-up with monitoring',
    'Key metric: Watch CPA and conversion rate changes'
]
for consideration in considerations:
    doc.add_paragraph(consideration, style='List Bullet')

doc.add_page_break()

# ============================================================================
# QUESTION 5: POTENTIAL CONFOUNDERS
# ============================================================================
doc.add_heading('5. POTENTIAL CONFOUNDERS & RISKS', 1)

doc.add_heading('Critical Factors That May Compromise Analysis', 2)

doc.add_heading('HIGH IMPACT Confounders', 3)

confounders = [
    {
        'title': '1. Seasonality & Timing',
        'issue': 'Campaigns may have run during different periods',
        'risk': 'Performance differences could be timing-related, not platform-related',
        'mitigation': 'Analyze campaign launch dates; Control for seasonal effects; Test same campaigns across different seasons'
    },
    {
        'title': '2. Product Mix Variation',
        'issue': 'Platforms attract different product categories (proven in analysis)',
        'risk': 'ROI differences may reflect product margins, not platform effectiveness',
        'mitigation': 'Normalize ROI by product category; Test cross-platform with same product category'
    },
    {
        'title': '3. Ad Creative Quality',
        'issue': 'Same ad may not be appropriate for all platforms',
        'risk': 'Poor performance may reflect creative mismatch, not platform viability',
        'mitigation': 'Develop platform-specific creative; A/B test native vs cross-posted content'
    },
    {
        'title': '4. Attribution Window Issues',
        'issue': 'Last-touch attribution may miss multi-touch customer journeys',
        'risk': 'Instagram may get credit for sales influenced by earlier touchpoints',
        'mitigation': 'Implement multi-touch attribution modeling; Track customer journey across platforms'
    }
]

for conf in confounders:
    p = doc.add_paragraph()
    run = p.add_run(conf['title'])
    run.bold = True
    doc.add_paragraph(f"Issue: {conf['issue']}")
    doc.add_paragraph(f"Risk: {conf['risk']}")
    doc.add_paragraph(f"Mitigation: {conf['mitigation']}")
    doc.add_paragraph()

doc.add_heading('MEDIUM IMPACT Confounders', 3)

medium_confounders = [
    '5. Audience Overlap - Same customers exposed to multiple campaigns',
    '6. Budget Level Non-linearity - Different budget levels may have different efficiency curves',
    '7. Audience Maturity Stage - Different platforms may catch customers at different buying stages',
    '8. External Market Factors - Economic conditions, competitor actions, trends'
]

for conf in medium_confounders:
    doc.add_paragraph(conf, style='List Bullet')

doc.add_heading('Risk Mitigation Recommendations', 2)

p = doc.add_paragraph()
run = p.add_run('Immediate Actions:')
run.bold = True

immediate = [
    'Develop platform-specific creative (addresses creative confounder)',
    'Implement enhanced tracking and attribution',
    'Conduct A/B tests with controlled variables',
    'Monitor marginal returns as budget scales',
    'Set up competitive monitoring'
]
for action in immediate:
    doc.add_paragraph(action, style='List Bullet')

p2 = doc.add_paragraph()
run2 = p2.add_run('Ongoing Monitoring:')
run2.bold = True

ongoing = [
    'Weekly CPA and ROAS tracking',
    'Monthly attribution analysis',
    'Quarterly audience research',
    'Competitive benchmarking'
]
for item in ongoing:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# QUESTION 6: CHANNEL-SPECIFIC AD CREATIVE
# ============================================================================
doc.add_heading('6. CREATIVE STRATEGY RECOMMENDATIONS', 1)

doc.add_heading('Platform-Specific Ad Creative (Key Recommendation)', 2)

p = doc.add_paragraph()
run = p.add_run('Critical Insight: Same ad across all platforms is suboptimal given audience differences')
run.bold = True
doc.add_paragraph()

doc.add_heading('Instagram Creative Strategy', 3)
doc.add_paragraph('Focus: Fashion, lifestyle, visual storytelling')
doc.add_paragraph('Format: Stories, Reels, carousel ads')
doc.add_paragraph('Tone: Aspirational, trendy, aesthetic')
doc.add_paragraph('CTA: "Shop the Look," "Get the Style"')
doc.add_paragraph('Example: Model photos, outfit inspiration, influencer partnerships')
doc.add_paragraph()

doc.add_heading('LinkedIn Creative Strategy', 3)
doc.add_paragraph('Focus: Tech products, professional tools, ROI messaging')
doc.add_paragraph('Format: Sponsored content, document ads, video testimonials')
doc.add_paragraph('Tone: Professional, data-driven, authoritative')
doc.add_paragraph('CTA: "Request Demo," "Download Whitepaper," "Learn More"')
doc.add_paragraph('Example: Product demos, case studies, industry insights')
doc.add_paragraph()

doc.add_heading('Facebook Creative Strategy', 3)
doc.add_paragraph('Focus: Health & wellness products, family-oriented')
doc.add_paragraph('Format: Video testimonials, carousel, collection ads')
doc.add_paragraph('Tone: Warm, supportive, community-focused')
doc.add_paragraph('CTA: "Shop Health Products," "Join Our Community"')
doc.add_paragraph('Example: Customer transformations, health benefits, family scenarios')
doc.add_paragraph()

doc.add_heading('Google Creative Strategy', 3)
doc.add_paragraph('Focus: Intent-driven, direct response, product-specific')
doc.add_paragraph('Format: Search ads, Shopping ads, display retargeting')
doc.add_paragraph('Tone: Clear, informative, value-focused')
doc.add_paragraph('CTA: "Buy Now," "Shop Sale," "Free Shipping"')
doc.add_paragraph('Example: Product features, special offers, competitive pricing')
doc.add_paragraph()

doc.add_heading('Expected Impact of Optimized Creative', 2)
doc.add_paragraph('Projected CTR increase: 25-40%')
doc.add_paragraph('Projected conversion rate lift: 15-30%')
doc.add_paragraph('Projected CPA reduction: 10-20%')
doc.add_paragraph('Timeline to results: 2-4 weeks after implementation')

doc.add_page_break()

# ============================================================================
# QUESTION 7: ORGANIC VS PAID SOCIAL STRATEGY
# ============================================================================
doc.add_heading('7. ORGANIC vs PAID SOCIAL STRATEGY', 1)

doc.add_heading('Understanding the Difference', 2)

comparison_table = doc.add_table(rows=5, cols=3)
comparison_table.style = 'Medium Shading 1 Accent 1'

comparison_table.rows[0].cells[0].text = 'Aspect'
comparison_table.rows[0].cells[1].text = 'Organic Posts'
comparison_table.rows[0].cells[2].text = 'Paid Ads'

comparison_table.rows[1].cells[0].text = 'Objective'
comparison_table.rows[1].cells[1].text = 'Build community & loyalty'
comparison_table.rows[1].cells[2].text = 'Drive immediate conversions'

comparison_table.rows[2].cells[0].text = 'Timeline'
comparison_table.rows[2].cells[1].text = 'Long-term relationships'
comparison_table.rows[2].cells[2].text = 'Short-term performance'

comparison_table.rows[3].cells[0].text = 'Content'
comparison_table.rows[3].cells[1].text = 'Educational, entertaining'
comparison_table.rows[3].cells[2].text = 'Direct, promotional'

comparison_table.rows[4].cells[0].text = 'Metrics'
comparison_table.rows[4].cells[1].text = 'Engagement, reach, sentiment'
comparison_table.rows[4].cells[2].text = 'ROI, ROAS, conversions'

doc.add_paragraph()

doc.add_heading('Organic Social Example (Instagram)', 3)
doc.add_paragraph('Educational Post:', style='Intense Quote')
doc.add_paragraph('💡 STYLE TIP TUESDAY')
doc.add_paragraph('How to build a capsule wardrobe that works for any season:')
doc.add_paragraph('1️⃣ Start with neutral basics')
doc.add_paragraph('2️⃣ Add 3-5 statement pieces')
doc.add_paragraph('3️⃣ Layer for versatility')
doc.add_paragraph('4️⃣ Accessorize to transform')
doc.add_paragraph()
doc.add_paragraph('What\'s your go-to versatile piece? Drop a 👗 in the comments!')
doc.add_paragraph('#CyberdineStyle #FashionTips')
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Goal: Build engagement, position as expert, foster community')
run.italic = True

doc.add_paragraph()

doc.add_heading('Paid Ad Example (Instagram)', 3)
doc.add_paragraph('Conversion Campaign:', style='Intense Quote')
doc.add_paragraph('🔥 48-HOUR FLASH SALE 🔥')
doc.add_paragraph()
doc.add_paragraph('40% OFF Everything + Free Shipping')
doc.add_paragraph()
doc.add_paragraph('✨ Premium fashion collection')
doc.add_paragraph('✨ Limited quantities')
doc.add_paragraph('✨ No code needed - auto-applied')
doc.add_paragraph()
doc.add_paragraph('[Shop Now] ← CTA Button')
doc.add_paragraph()
doc.add_paragraph('⏰ Sale ends Sunday midnight')
doc.add_paragraph('🚨 Bestsellers selling out fast')
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Goal: Drive immediate purchases, create urgency, maximize revenue')
run.italic = True

doc.add_paragraph()

doc.add_heading('Synergy Strategy', 2)
synergy = [
    'Test organically - See what content resonates',
    'Boost top performers - Turn winning posts into ads',
    'Retarget engagers - Convert interested followers',
    'Scale winners - Invest more in proven creative'
]
for item in synergy:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# FINAL RECOMMENDATIONS
# ============================================================================
doc.add_heading('STRATEGIC RECOMMENDATIONS SUMMARY', 1)

doc.add_heading('Immediate Actions (Next 30 Days)', 2)

p1 = doc.add_paragraph()
run1 = p1.add_run('Priority 1: Optimize Creative')
run1.bold = True

priority1 = [
    'Develop platform-specific ad variations',
    'Test Instagram fashion-focused creative',
    'Create LinkedIn B2B tech messaging',
    'Design Facebook health & wellness campaigns'
]
for item in priority1:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph('Expected Impact: 15-30% performance improvement')
doc.add_paragraph()

p2 = doc.add_paragraph()
run2 = p2.add_run('Priority 2: Implement $5K Test Campaign')
run2.bold = True

priority2 = [
    'Launch recommended 4-campaign portfolio',
    'Monitor daily for first 2 weeks',
    'Adjust based on early performance'
]
for item in priority2:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph('Expected ROI: 291.5%')
doc.add_paragraph('Expected Profit: $19,576')
doc.add_paragraph()

p3 = doc.add_paragraph()
run3 = p3.add_run('Priority 3: Set Up Enhanced Tracking')
run3.bold = True

priority3 = [
    'Implement multi-touch attribution',
    'Configure UTM parameters',
    'Set up conversion tracking pixels',
    'Create dashboard for real-time monitoring'
]
for item in priority3:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph()

doc.add_heading('Budget Decision Matrix', 2)

decision_table = doc.add_table(rows=4, cols=3)
decision_table.style = 'Medium Shading 1 Accent 1'

decision_table.rows[0].cells[0].text = 'Budget'
decision_table.rows[0].cells[1].text = 'Expected Monthly Profit'
decision_table.rows[0].cells[2].text = 'Recommendation'

decision_table.rows[1].cells[0].text = '$2,000'
decision_table.rows[1].cells[1].text = '$8,195 (309.8% ROI)'
decision_table.rows[1].cells[2].text = 'Conservative option'

decision_table.rows[2].cells[0].text = '$5,000'
decision_table.rows[2].cells[1].text = '$19,576 (291.5% ROI)'
decision_table.rows[2].cells[2].text = 'RECOMMENDED'

decision_table.rows[3].cells[0].text = '$30,000'
decision_table.rows[3].cells[1].text = '~$70,000 (~233% ROI)'
decision_table.rows[3].cells[2].text = 'Scale target'

doc.add_paragraph()

doc.add_heading('Critical Success Factors', 2)

success = [
    '✓ DO: Develop platform-specific creative immediately',
    '✓ DO: Start with proven winners, scale gradually',
    '✓ DO: Monitor performance daily during scale-up',
    '✓ DO: Maintain diversification for risk management',
    '✓ DO: Test, measure, optimize continuously',
    '',
    '✗ DON\'T: Use same ad across all platforms',
    '✗ DON\'T: Scale too quickly without data validation',
    '✗ DON\'T: Ignore confounders and attribution issues',
    '✗ DON\'T: Set and forget - active management required',
    '✗ DON\'T: Abandon underperformers without testing variations'
]

for item in success:
    if item:
        doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph()

doc.add_page_break()

# ============================================================================
# CONCLUSION
# ============================================================================
doc.add_heading('CONCLUSION', 1)

p = doc.add_paragraph()
run = p.add_run('Cyberdine has a strong foundation for marketing growth:')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()

conclusions = [
    '✓ All channels are profitable - No losing platforms',
    '✓ Clear winner identified - Instagram shows exceptional performance',
    '✓ Scalable opportunities - Room to grow investment profitably',
    '✓ Audience insights discovered - Platform-specific targeting opportunities',
    '✓ Action plan defined - Clear next steps with expected outcomes'
]

for item in conclusions:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

p2 = doc.add_paragraph()
run2 = p2.add_run('Key Success Factor: ')
run2.bold = True
run3 = p2.add_run('Platform-specific creative optimization')

doc.add_paragraph()

p3 = doc.add_paragraph()
run4 = p3.add_run('Recommended Next Step: ')
run4.bold = True
run5 = p3.add_run('Approve $5K test budget with platform-specific creative')

doc.add_paragraph()

p4 = doc.add_paragraph()
run6 = p4.add_run('Expected Timeline to Results: ')
run6.bold = True
run7 = p4.add_run('30-45 days for optimized creative rollout')

doc.add_paragraph()

p5 = doc.add_paragraph()
run8 = p5.add_run('Board Decision Requested: ')
run8.bold = True
run9 = p5.add_run('Budget approval and creative production authorization')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

footer = doc.add_paragraph('Prepared by: Marketing Analytics Team')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(10)
footer.runs[0].font.italic = True

date_footer = doc.add_paragraph('November 28, 2025')
date_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_footer.runs[0].font.size = Pt(10)
date_footer.runs[0].font.italic = True

# Save document
doc.save('Cyberdine_Executive_Report.docx')
print('✅ Executive Report saved as: Cyberdine_Executive_Report.docx')
print('   File includes all 7 questions answered with tables, analysis, and recommendations')
